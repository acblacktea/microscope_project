import sys
import os
import io
import base64
import re

# Add uvcham SDK to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'uvchamsdk.20250428', 'python'))
import uvcham

from PyQt6.QtCore import (pyqtSignal, Qt, QTimer, QThread, QBuffer,
                          QIODevice, QPropertyAnimation, QEasingCurve, QPoint)
from PyQt6.QtGui import QPixmap, QImage, QFont
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QLabel, QPushButton,
    QVBoxLayout, QHBoxLayout, QTextEdit, QMessageBox, QMenu,
    QSplitter, QFrame, QSizePolicy, QScrollArea, QGridLayout,
    QFileDialog, QSlider, QTabWidget
)
from PyQt6.QtPrintSupport import QPrinter

from gemini_service import analyze_images


class CameraWidget(QWidget):
    """左侧面板：显微镜实时视频画面"""
    evtCallback = pyqtSignal(int)

    HOVER_ZONE_WIDTH = 40  # 鼠标靠近左边缘多少像素时触发

    def __init__(self, parent=None):
        super().__init__(parent)
        self.hcam = None
        self.imgWidth = 0
        self.imgHeight = 0
        self.pData = None
        self.frame = 0
        self.timer = QTimer(self)
        self.setMouseTracking(True)

        # 视频显示区域
        self.lbl_video = QLabel("未连接相机")
        self.lbl_video.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.lbl_video.setStyleSheet(
            "QLabel { background-color: #1a1a2e; color: #888; "
            "font-size: 16px; border: 1px solid #333; border-radius: 4px; }"
        )
        self.lbl_video.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self.lbl_video.setMouseTracking(True)

        # 相机控制栏
        self.btn_open = QPushButton("打开相机")
        self.btn_open.setFixedHeight(36)
        self.btn_open.clicked.connect(self.onBtnOpen)

        self.lbl_status = QLabel("状态：未连接")
        self.lbl_status.setStyleSheet("color: #888; font-size: 12px;")

        ctrl_layout = QHBoxLayout()
        ctrl_layout.addWidget(self.btn_open)
        ctrl_layout.addStretch()
        ctrl_layout.addWidget(self.lbl_status)

        layout = QVBoxLayout()
        layout.addWidget(self.lbl_video, 1)
        layout.addLayout(ctrl_layout)
        self.setLayout(layout)

        self.evtCallback.connect(self.onEvtCallback)
        self.timer.timeout.connect(self.onTimer)

        # 浮动调节面板（作为子 widget 覆盖在画面左侧）
        self.adjust_overlay = None  # 延迟初始化，等 layout 完成
        self._overlay_visible = False
        self._hide_timer = QTimer(self)
        self._hide_timer.setSingleShot(True)
        self._hide_timer.setInterval(400)
        self._hide_timer.timeout.connect(self._hideOverlay)

    def _ensureOverlay(self):
        """延迟创建浮动调节面板"""
        if self.adjust_overlay is not None:
            return
        self.adjust_overlay = ImageAdjustPanel(self, parent=self)
        self.adjust_overlay.setFixedWidth(220)
        self.adjust_overlay.setMouseTracking(True)
        self.adjust_overlay.setAutoFillBackground(True)
        self.adjust_overlay.setAttribute(Qt.WidgetAttribute.WA_StyledBackground, True)
        self.adjust_overlay.setObjectName("adjustOverlay")
        self.adjust_overlay.setStyleSheet("""
            QWidget#adjustOverlay {
                background-color: rgba(22, 33, 62, 230);
                border: 1px solid #4a9eff;
                border-radius: 8px;
            }
            QWidget#adjustOverlay QLabel {
                background: transparent;
            }
            QWidget#adjustOverlay QPushButton {
                background-color: #2a9d8f;
            }
            QWidget#adjustOverlay QWidget {
                background: transparent;
            }
        """)
        # 初始位置在可视区域外（左侧隐藏）
        self.adjust_overlay.move(-220, 40)
        self.adjust_overlay.show()

    def _showOverlay(self):
        """滑入浮动面板"""
        self._ensureOverlay()
        if self._overlay_visible:
            return
        self._hide_timer.stop()
        self._overlay_visible = True
        self.adjust_overlay.setFixedHeight(self.height() - 80)
        self.adjust_overlay.raise_()  # 提到最前面，避免被视频画面遮挡
        self.adjust_overlay.syncFromCamera()

        anim = QPropertyAnimation(self.adjust_overlay, b"pos", self)
        anim.setDuration(200)
        anim.setStartValue(self.adjust_overlay.pos())
        anim.setEndValue(QPoint(8, 40))
        anim.setEasingCurve(QEasingCurve.Type.OutCubic)
        anim.start()
        self._anim = anim  # prevent GC

    def _hideOverlay(self):
        """滑出浮动面板"""
        if not self._overlay_visible or self.adjust_overlay is None:
            return
        self._overlay_visible = False

        anim = QPropertyAnimation(self.adjust_overlay, b"pos", self)
        anim.setDuration(200)
        anim.setStartValue(self.adjust_overlay.pos())
        anim.setEndValue(QPoint(-220, 40))
        anim.setEasingCurve(QEasingCurve.Type.InCubic)
        anim.start()
        self._anim = anim

    def mouseMoveEvent(self, event):
        """鼠标在画面左边缘区域时显示调节面板"""
        pos = event.position() if hasattr(event, 'position') else event.pos()
        x = int(pos.x())
        if x <= self.HOVER_ZONE_WIDTH:
            self._showOverlay()
        elif self._overlay_visible and self.adjust_overlay is not None:
            # 检查鼠标是否在浮动面板区域内
            overlay_rect = self.adjust_overlay.geometry()
            mouse_pt = QPoint(int(pos.x()), int(pos.y()))
            if not overlay_rect.contains(mouse_pt):
                self._hide_timer.start()
        super().mouseMoveEvent(event)

    def leaveEvent(self, event):
        """鼠标离开 CameraWidget 时延迟隐藏"""
        if self._overlay_visible:
            self._hide_timer.start()
        super().leaveEvent(event)

    @staticmethod
    def cameraCallback(nEvent, ctx):
        ctx.evtCallback.emit(nEvent)

    def onEvtCallback(self, nEvent):
        if self.hcam is not None:
            if uvcham.UVCHAM_EVENT_IMAGE & nEvent != 0:
                self.onImageEvent()
            elif uvcham.UVCHAM_EVENT_ERROR & nEvent != 0:
                self.closeCamera()
                QMessageBox.warning(self, "警告", "相机发生错误。")
            elif uvcham.UVCHAM_EVENT_DISCONNECT & nEvent != 0:
                self.closeCamera()
                QMessageBox.warning(self, "警告", "相机已断开连接。")

    def onImageEvent(self):
        self.hcam.pull(self.pData)
        self.frame += 1
        image = QImage(self.pData, self.imgWidth, self.imgHeight, QImage.Format.Format_RGB888)
        scaled = image.scaled(
            self.lbl_video.width(), self.lbl_video.height(),
            Qt.AspectRatioMode.KeepAspectRatio,
            Qt.TransformationMode.FastTransformation
        )
        self.lbl_video.setPixmap(QPixmap.fromImage(scaled))

    def onTimer(self):
        if self.hcam is not None:
            self.lbl_status.setText(f"状态：运行中 | 帧数：{self.frame}")

    def openCamera(self, cam_id):
        self.hcam = uvcham.Uvcham.open(cam_id)
        if self.hcam:
            self.frame = 0
            self.hcam.put(uvcham.UVCHAM_FORMAT, 2)  # RGB888 for QImage

            res = self.hcam.get(uvcham.UVCHAM_RES)
            self.imgWidth = self.hcam.get(uvcham.UVCHAM_WIDTH | res)
            self.imgHeight = self.hcam.get(uvcham.UVCHAM_HEIGHT | res)
            self.pData = bytes(uvcham.TDIBWIDTHBYTES(self.imgWidth * 24) * self.imgHeight)

            try:
                self.hcam.start(None, self.cameraCallback, self)  # Pull Mode
            except uvcham.HRESULTException:
                self.closeCamera()
                QMessageBox.warning(self, "警告", "启动相机失败。")
            else:
                self.btn_open.setText("关闭相机")
                self.lbl_status.setText("状态：运行中 | 帧数：0")
                self.timer.start(1000)

    def onBtnOpen(self):
        if self.hcam is not None:
            self.closeCamera()
        else:
            arr = uvcham.Uvcham.enum()
            if len(arr) == 0:
                QMessageBox.warning(self, "警告", "未找到相机。")
            elif len(arr) == 1:
                self.openCamera(arr[0].id)
            else:
                menu = QMenu(self)
                for i, dev in enumerate(arr):
                    action = menu.addAction(dev.displayname)
                    action.setData(i)
                action = menu.exec(self.mapToGlobal(self.btn_open.pos()))
                if action:
                    self.openCamera(arr[action.data()].id)

    def closeCamera(self):
        if self.hcam:
            self.hcam.close()
        self.hcam = None
        self.pData = None
        self.btn_open.setText("打开相机")
        self.lbl_status.setText("状态：未连接")
        self.lbl_video.clear()
        self.lbl_video.setText("未连接相机")
        self.timer.stop()

    def getCurrentFrame(self):
        """返回当前帧的 QImage，无相机时返回 None。"""
        if self.hcam is not None and self.pData is not None:
            return QImage(self.pData, self.imgWidth, self.imgHeight, QImage.Format.Format_RGB888)
        return None


class GeminiWorker(QThread):
    """后台线程：调用 Gemini API 避免阻塞 UI"""
    finished = pyqtSignal(str)

    def __init__(self, image_data_list: list, mode: str = "algae", parent=None):
        super().__init__(parent)
        self.image_data_list = image_data_list
        self.mode = mode

    def run(self):
        try:
            result = analyze_images(self.image_data_list, mode=self.mode)
            self.finished.emit(result)
        except Exception as e:
            self.finished.emit(f"分析出错：{str(e)}")


class ThumbnailWidget(QWidget):
    """单张截图缩略图，右上角带删除按钮"""
    removed = pyqtSignal(int)

    def __init__(self, index: int, pixmap: QPixmap, parent=None):
        super().__init__(parent)
        self.index = index
        self.setFixedSize(90, 90)

        # 缩略图
        lbl_img = QLabel(self)
        lbl_img.setPixmap(pixmap.scaled(
            84, 84,
            Qt.AspectRatioMode.KeepAspectRatio,
            Qt.TransformationMode.SmoothTransformation
        ))
        lbl_img.setAlignment(Qt.AlignmentFlag.AlignCenter)
        lbl_img.setGeometry(0, 0, 90, 90)
        lbl_img.setStyleSheet("background-color: #1a1a2e; border: 1px solid #444; border-radius: 4px;")

        # 右上角删除按钮
        btn_del = QPushButton("-", self)
        btn_del.setFixedSize(20, 20)
        btn_del.move(66, 2)
        btn_del.setStyleSheet("""
            QPushButton {
                background-color: #e04040; color: white; font-weight: bold;
                border: none; border-radius: 10px; font-size: 14px; padding: 0;
            }
            QPushButton:hover { background-color: #ff5555; }
        """)
        btn_del.clicked.connect(lambda: self.removed.emit(self.index))


class AnalysisPanel(QWidget):
    """右侧面板：截图管理 + AI 分析结果展示"""

    # 模式配置
    MODE_CONFIG = {
        "algae": {
            "title": "藻类智能分析",
            "report_title": "藻类AI智慧分析报告",
            "export_name": "藻类分析报告",
            "capture_count": 2,
        },
        "shrimp": {
            "title": "虾体健康分析",
            "report_title": "虾体AI健康诊断报告",
            "export_name": "虾体健康报告",
            "capture_count": 1,
        },
    }

    def __init__(self, camera_widget: CameraWidget, mode: str = "algae", parent=None):
        super().__init__(parent)
        self.camera_widget = camera_widget
        self.mode = mode
        self.cfg = self.MODE_CONFIG.get(mode, self.MODE_CONFIG["algae"])
        self.captured_images = []   # PNG bytes list
        self._raw_result = ""       # AI 返回的原始 HTML
        self.capture_count = 0      # 当前连拍已截取张数
        self.capture_timer = QTimer(self)
        self.capture_timer.timeout.connect(self._captureOneFrame)
        self.worker = None

        # 标题
        title = QLabel(self.cfg["title"])
        title.setFont(QFont("Arial", 14, QFont.Weight.Bold))
        title.setStyleSheet("color: #e0e0e0; padding: 4px 0;")

        # 截取图像按钮
        self.btn_capture = QPushButton("截取图像")
        self.btn_capture.setFixedHeight(38)
        self.btn_capture.setStyleSheet("""
            QPushButton {
                background-color: #2a9d8f; color: white; font-size: 14px;
                font-weight: bold; border: none; border-radius: 6px;
            }
            QPushButton:hover { background-color: #21867a; }
            QPushButton:pressed { background-color: #1a6f66; }
        """)
        self.btn_capture.clicked.connect(self.onCapture)

        # 缩略图展示区域（用 FlowLayout 模拟：ScrollArea + GridLayout）
        self.thumb_container = QWidget()
        self.thumb_layout = QGridLayout(self.thumb_container)
        self.thumb_layout.setContentsMargins(0, 0, 0, 0)
        self.thumb_layout.setSpacing(6)
        self.thumb_container.setStyleSheet("background: transparent;")

        self.thumb_scroll = QScrollArea()
        self.thumb_scroll.setWidget(self.thumb_container)
        self.thumb_scroll.setWidgetResizable(True)
        self.thumb_scroll.setFixedHeight(110)
        self.thumb_scroll.setStyleSheet("""
            QScrollArea { background: transparent; border: none; }
            QScrollBar:vertical { width: 6px; background: transparent; }
            QScrollBar::handle:vertical { background: #555; border-radius: 3px; }
        """)

        # 分析按钮 + 导出PDF按钮
        self.btn_analyze = QPushButton("分析当前画面")
        self.btn_analyze.setFixedHeight(44)
        self.btn_analyze.setStyleSheet("""
            QPushButton {
                background-color: #4a9eff; color: white; font-size: 15px;
                font-weight: bold; border: none; border-radius: 6px;
            }
            QPushButton:hover { background-color: #3a8eef; }
            QPushButton:pressed { background-color: #2a7edf; }
            QPushButton:disabled { background-color: #555; color: #999; }
        """)
        self.btn_analyze.clicked.connect(self.onAnalyze)

        self.btn_export = QPushButton("导出PDF")
        self.btn_export.setFixedHeight(44)
        self.btn_export.setStyleSheet("""
            QPushButton {
                background-color: #e07a2f; color: white; font-size: 15px;
                font-weight: bold; border: none; border-radius: 6px;
            }
            QPushButton:hover { background-color: #c96a22; }
            QPushButton:pressed { background-color: #b05a18; }
            QPushButton:disabled { background-color: #555; color: #999; }
        """)
        self.btn_export.clicked.connect(self.onExportPDF)

        self.btn_export_word = QPushButton("导出Word")
        self.btn_export_word.setFixedHeight(44)
        self.btn_export_word.setStyleSheet("""
            QPushButton {
                background-color: #2b5797; color: white; font-size: 15px;
                font-weight: bold; border: none; border-radius: 6px;
            }
            QPushButton:hover { background-color: #234a82; }
            QPushButton:pressed { background-color: #1c3d6e; }
            QPushButton:disabled { background-color: #555; color: #999; }
        """)
        self.btn_export_word.clicked.connect(self.onExportWord)

        btn_row = QHBoxLayout()
        btn_row.setSpacing(8)
        btn_row.addWidget(self.btn_analyze, 3)
        btn_row.addWidget(self.btn_export_word, 1)
        btn_row.addWidget(self.btn_export, 1)

        # 状态标签
        self.lbl_status = QLabel("")
        self.lbl_status.setStyleSheet("color: #4a9eff; font-size: 12px; background: transparent;")

        # 结果显示区域
        self.txt_results = QTextEdit()
        self.txt_results.setReadOnly(True)
        self.txt_results.setPlaceholderText("分析结果将在此处显示...")
        self.txt_results.setStyleSheet("""
            QTextEdit {
                background-color: #1a1a2e; color: #d0d0d0;
                border: 1px solid #333; border-radius: 4px;
                padding: 8px; font-size: 13px;
            }
        """)

        layout = QVBoxLayout()
        layout.addWidget(title)
        layout.addWidget(self.btn_capture)
        layout.addWidget(self.thumb_scroll)
        layout.addLayout(btn_row)
        layout.addWidget(self.lbl_status)
        layout.addWidget(self.txt_results, 1)
        self.setLayout(layout)

    def onCapture(self):
        """点击截取：根据模式每隔1秒拍指定张数"""
        frame = self.camera_widget.getCurrentFrame()
        if frame is None:
            self.lbl_status.setText("请先打开相机。")
            return

        self._capture_total = self.cfg["capture_count"]
        self.btn_capture.setEnabled(False)
        self.capture_count = 0
        # 立即拍第一张
        self._captureOneFrame()
        # 如果需要多张，后续每隔1秒拍一张
        if self._capture_total > 1:
            self.capture_timer.start(1000)

    def _captureOneFrame(self):
        """采集一帧并转为 PNG 字节数据"""
        frame = self.camera_widget.getCurrentFrame()
        if frame is not None:
            buf = QBuffer()
            buf.open(QIODevice.OpenModeFlag.WriteOnly)
            frame.save(buf, "PNG")
            self.captured_images.append(bytes(buf.data()))
            buf.close()

        self.capture_count += 1
        total = self._capture_total
        self.lbl_status.setText(f"正在截取... ({self.capture_count}/{total})")
        self.refreshThumbnails()

        if self.capture_count >= total:
            self.capture_timer.stop()
            self.btn_capture.setEnabled(True)
            self.lbl_status.setText(f"已截取 {len(self.captured_images)} 张图像")

    def removeImage(self, index: int):
        """删除指定索引的截图"""
        if 0 <= index < len(self.captured_images):
            self.captured_images.pop(index)
            self.refreshThumbnails()
            self.lbl_status.setText(f"已截取 {len(self.captured_images)} 张图像")

    def refreshThumbnails(self):
        """刷新缩略图展示"""
        # 清除旧的缩略图
        while self.thumb_layout.count():
            item = self.thumb_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        cols = 3
        for i, img_data in enumerate(self.captured_images):
            pixmap = QPixmap()
            pixmap.loadFromData(img_data, "PNG")
            thumb = ThumbnailWidget(i, pixmap)
            thumb.removed.connect(self.removeImage)
            self.thumb_layout.addWidget(thumb, i // cols, i % cols)

    def onExportWord(self):
        """将分析结果导出为 Word 文档"""
        if not self._raw_result:
            self.lbl_status.setText("没有可导出的分析结果。")
            return

        file_path, _ = QFileDialog.getSaveFileName(
            self, "导出Word", f"{self.cfg['export_name']}.docx", "Word文件 (*.docx)"
        )
        if not file_path:
            return

        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except ImportError:
            self.lbl_status.setText("缺少 python-docx，请执行 pip install python-docx")
            return

        FONT_NAME = '微软雅黑'

        doc = Document()

        # 设置默认字体为中文字体
        style = doc.styles['Normal']
        style.font.name = FONT_NAME
        style.font.size = Pt(11)
        style.element.rPr.rFonts.set(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', FONT_NAME
        )

        # 标题
        title = doc.add_heading(self.cfg['report_title'], level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.name = FONT_NAME
            run._element.rPr.rFonts.set(
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', FONT_NAME
            )

        def set_run_font(run, font_name=FONT_NAME, size=None, bold=False, color=None):
            """给 run 设置中文字体"""
            run.font.name = font_name
            run._element.rPr.rFonts.set(
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', font_name
            )
            if size:
                run.font.size = size
            if bold:
                run.bold = True
            if color:
                run.font.color.rgb = color

        def set_paragraph_font(paragraph):
            """给段落所有 run 设置中文字体"""
            for run in paragraph.runs:
                set_run_font(run)

        # 解析 HTML 内容写入 Word
        html = self._raw_result
        # 去除可能的 ```html 包裹
        html = re.sub(r'^```html\s*', '', html.strip())
        html = re.sub(r'\s*```$', '', html.strip())

        # 简单 HTML 解析：按标签拆分内容
        parts = re.split(r'(<h2>.*?</h2>|<table>.*?</table>|<ul>.*?</ul>|<li>.*?</li>)', html, flags=re.DOTALL)

        in_list = False
        for part in parts:
            part = part.strip()
            if not part:
                continue

            # h2 标题
            if part.startswith('<h2>'):
                text = re.sub(r'<.*?>', '', part)
                h = doc.add_heading(text, level=2)
                for run in h.runs:
                    set_run_font(run)
                in_list = False

            # 表格
            elif part.startswith('<table>'):
                rows_html = re.findall(r'<tr>(.*?)</tr>', part, re.DOTALL)
                if rows_html:
                    # 解析第一行确定列数
                    first_cells = re.findall(r'<t[hd]>(.*?)</t[hd]>', rows_html[0], re.DOTALL)
                    num_cols = len(first_cells)
                    if num_cols > 0:
                        table = doc.add_table(rows=0, cols=num_cols)
                        table.style = 'Table Grid'
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER

                        for row_idx, row_html in enumerate(rows_html):
                            cells_text = re.findall(r'<t[hd]>(.*?)</t[hd]>', row_html, re.DOTALL)
                            row = table.add_row()
                            for col_idx, cell_text in enumerate(cells_text):
                                clean = re.sub(r'<.*?>', '', cell_text).strip()
                                cell = row.cells[col_idx]
                                cell.text = clean
                                for p in cell.paragraphs:
                                    for run in p.runs:
                                        set_run_font(run,
                                                     bold=(row_idx == 0),
                                                     color=RGBColor(0x1a, 0x5f, 0xb4) if row_idx == 0 else None)
                        doc.add_paragraph()  # 表格后空行
                in_list = False

            # 列表项
            elif part.startswith('<li>'):
                text = re.sub(r'<.*?>', '', part).strip()
                if text:
                    p = doc.add_paragraph(text, style='List Bullet')
                    set_paragraph_font(p)
                in_list = True

            # ul 块（里面的 li 可能没被单独匹配到）
            elif part.startswith('<ul>'):
                items = re.findall(r'<li>(.*?)</li>', part, re.DOTALL)
                for item in items:
                    text = re.sub(r'<.*?>', '', item).strip()
                    if text:
                        p = doc.add_paragraph(text, style='List Bullet')
                        set_paragraph_font(p)
                in_list = True

            # 普通文本
            else:
                text = re.sub(r'<.*?>', '', part).strip()
                if text and not in_list:
                    p = doc.add_paragraph(text)
                    set_paragraph_font(p)

        # 添加截取的图片
        if self.captured_images:
            h = doc.add_heading('显微镜样本图像', level=2)
            for run in h.runs:
                set_run_font(run)
            for img_data in self.captured_images:
                doc.add_picture(io.BytesIO(img_data), width=Inches(2.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save(file_path)
        self.lbl_status.setText(f"已导出: {os.path.basename(file_path)}")

    def onExportPDF(self):
        """将分析结果导出为 PDF"""
        html = self.txt_results.toHtml()
        if not html.strip() or self.txt_results.toPlainText().strip() == "":
            self.lbl_status.setText("没有可导出的分析结果。")
            return

        file_path, _ = QFileDialog.getSaveFileName(
            self, "导出PDF", f"{self.cfg['export_name']}.pdf", "PDF文件 (*.pdf)"
        )
        if not file_path:
            return

        # 构建图片 HTML
        images_html = '<h2>显微镜样本图像</h2><p>'
        for img_data in self.captured_images:
            b64 = base64.b64encode(img_data).decode('ascii')
            images_html += f'<img src="data:image/png;base64,{b64}" width="200" height="150" /> '
        images_html += '</p>'

        # 构建打印用 HTML（白底黑字，适合打印）
        print_html = f"""
        <style>
            body {{ font-family: 'Microsoft YaHei', 'SimSun', sans-serif; color: #222; }}
            h2 {{ color: #1a5fb4; font-size: 16px; margin: 14px 0 6px 0; padding-bottom: 3px; border-bottom: 1px solid #ccc; }}
            ul {{ margin: 4px 0; padding-left: 20px; }}
            li {{ margin: 3px 0; line-height: 1.6; }}
            table {{ width: 100%; border-collapse: collapse; margin: 8px 0; }}
            th {{ background-color: #e8f0fe; color: #1a5fb4; padding: 6px 8px; text-align: left; border: 1px solid #bbb; }}
            td {{ padding: 6px 8px; border: 1px solid #bbb; line-height: 1.5; }}
            tr:nth-child(even) {{ background-color: #f5f5f5; }}
        </style>
        <h1 style="text-align:center; color:#1a5fb4;">{self.cfg['report_title']}</h1>
        {self._raw_result}
        {images_html}
        """

        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
        printer.setOutputFileName(file_path)

        doc = self.txt_results.document().clone()
        doc.setHtml(print_html)
        page_rect = printer.pageRect(QPrinter.Unit.Point)
        doc.setPageSize(page_rect.size())
        doc.setDefaultFont(QFont("Microsoft YaHei", 11))
        doc.print(printer)

        self.lbl_status.setText(f"已导出: {os.path.basename(file_path)}")

    def onAnalyze(self):
        """将已截取的图像发送给 AI 分析"""
        if not self.captured_images:
            self.lbl_status.setText("请先截取至少一张图像。")
            return

        self.btn_analyze.setEnabled(False)
        self.btn_capture.setEnabled(False)
        self.txt_results.clear()
        self.lbl_status.setText("正在调用 AI 分析...")
        self.startGeminiAnalysis()

    def startGeminiAnalysis(self):
        """在后台线程中调用 Gemini API"""
        self.worker = GeminiWorker(self.captured_images, mode=self.mode)
        self.worker.finished.connect(self.onAnalysisFinished)
        self.worker.start()

    def onAnalysisFinished(self, result: str):
        """Gemini 返回结果后显示"""
        self._raw_result = result
        styled_html = f"""
        <style>
            body {{ font-family: 'Microsoft YaHei', sans-serif; color: #d0d0d0; }}
            h2 {{ color: #4a9eff; font-size: 15px; margin: 16px 0 8px 0; padding-bottom: 4px; border-bottom: 1px solid #333; }}
            ul {{ margin: 4px 0; padding-left: 20px; }}
            li {{ margin: 4px 0; line-height: 1.6; }}
            table {{ width: 100%; border-collapse: collapse; margin: 8px 0; }}
            th {{ background-color: #2a2a4a; color: #4a9eff; padding: 8px; text-align: left; border: 1px solid #333; }}
            td {{ padding: 8px; border: 1px solid #333; line-height: 1.5; }}
            tr:nth-child(even) {{ background-color: #1a1a3e; }}
        </style>
        {result}
        """
        self.txt_results.setHtml(styled_html)
        self.lbl_status.setText("分析完成")
        self.btn_analyze.setEnabled(True)
        self.btn_capture.setEnabled(True)
        self.worker = None


class ImageAdjustPanel(QWidget):
    """画面调节面板：饱和度、对比度、曝光度滑动条（浮动覆盖层）"""

    def __init__(self, camera_widget: CameraWidget, parent=None):
        super().__init__(parent)
        self.camera_widget = camera_widget
        self.setMouseTracking(True)

        layout = QVBoxLayout()
        layout.setContentsMargins(8, 12, 8, 12)
        layout.setSpacing(16)

        title = QLabel("画面调节")
        title.setFont(QFont("Arial", 16, QFont.Weight.Bold))
        title.setStyleSheet("color: #e0e0e0; padding: 8px 0;")
        layout.addWidget(title)

        # 定义三个参数：(名称, SDK常量)
        self.params = [
            ("饱和度", uvcham.UVCHAM_SATURATION),
            ("对比度", uvcham.UVCHAM_CONTRAST),
            ("曝光度", uvcham.UVCHAM_BRIGHTNESS),
        ]
        self.sliders = {}
        self.value_labels = {}

        for name, param_id in self.params:
            row_widget = self._createSliderRow(name, param_id)
            layout.addWidget(row_widget)

        # 重置按钮
        self.btn_reset = QPushButton("恢复默认值")
        self.btn_reset.setFixedHeight(38)
        self.btn_reset.setStyleSheet("""
            QPushButton {
                background-color: #2a9d8f; color: white; font-size: 14px;
                font-weight: bold; border: none; border-radius: 6px;
            }
            QPushButton:hover { background-color: #21867a; }
            QPushButton:pressed { background-color: #1a6f66; }
        """)
        self.btn_reset.clicked.connect(self.onReset)
        layout.addWidget(self.btn_reset)

        self.lbl_status = QLabel("")
        self.lbl_status.setStyleSheet("color: #4a9eff; font-size: 12px; background: transparent;")
        layout.addWidget(self.lbl_status)

        layout.addStretch()
        self.setLayout(layout)

    def _createSliderRow(self, name: str, param_id: int) -> QWidget:
        """创建一行：标签 + 滑动条 + 数值"""
        container = QWidget()
        v_layout = QVBoxLayout(container)
        v_layout.setContentsMargins(0, 0, 0, 0)
        v_layout.setSpacing(4)

        # 标签和数值
        h_layout = QHBoxLayout()
        lbl_name = QLabel(name)
        lbl_name.setStyleSheet("color: #e0e0e0; font-size: 14px; font-weight: bold; background: transparent;")
        lbl_value = QLabel("--")
        lbl_value.setStyleSheet("color: #4a9eff; font-size: 13px; background: transparent;")
        lbl_value.setAlignment(Qt.AlignmentFlag.AlignRight)
        h_layout.addWidget(lbl_name)
        h_layout.addStretch()
        h_layout.addWidget(lbl_value)
        v_layout.addLayout(h_layout)

        # 滑动条
        slider = QSlider(Qt.Orientation.Horizontal)
        slider.setStyleSheet("""
            QSlider::groove:horizontal {
                border: 1px solid #444; height: 6px; background: #1a1a2e; border-radius: 3px;
            }
            QSlider::handle:horizontal {
                background: #4a9eff; border: none; width: 16px; height: 16px;
                margin: -5px 0; border-radius: 8px;
            }
            QSlider::handle:horizontal:hover { background: #6ab4ff; }
            QSlider::sub-page:horizontal { background: #4a9eff; border-radius: 3px; }
        """)
        slider.valueChanged.connect(lambda val, p=param_id, l=lbl_value: self._onSliderChanged(p, val, l))
        v_layout.addWidget(slider)

        self.sliders[param_id] = slider
        self.value_labels[param_id] = lbl_value
        return container

    def _onSliderChanged(self, param_id: int, value: int, lbl: QLabel):
        """滑动条值变化时更新相机参数"""
        lbl.setText(str(value))
        hcam = self.camera_widget.hcam
        if hcam is not None:
            try:
                hcam.put(param_id, value)
            except Exception:
                pass

    def syncFromCamera(self):
        """从相机读取当前参数值和范围，同步到滑动条"""
        hcam = self.camera_widget.hcam
        if hcam is None:
            self.lbl_status.setText("请先打开相机。")
            return

        for name, param_id in self.params:
            try:
                min_val, max_val, default_val = hcam.range(param_id)
                cur_val = hcam.get(param_id)
                slider = self.sliders[param_id]
                slider.blockSignals(True)
                slider.setRange(min_val, max_val)
                slider.setValue(cur_val)
                slider.blockSignals(False)
                self.value_labels[param_id].setText(str(cur_val))
            except Exception:
                pass
        self.lbl_status.setText("已同步相机参数。")

    def onReset(self):
        """恢复默认值"""
        hcam = self.camera_widget.hcam
        if hcam is None:
            self.lbl_status.setText("请先打开相机。")
            return

        for name, param_id in self.params:
            try:
                min_val, max_val, default_val = hcam.range(param_id)
                hcam.put(param_id, default_val)
                slider = self.sliders[param_id]
                slider.blockSignals(True)
                slider.setValue(default_val)
                slider.blockSignals(False)
                self.value_labels[param_id].setText(str(default_val))
            except Exception:
                pass
        self.lbl_status.setText("已恢复默认值。")

    def enterEvent(self, event):
        """鼠标进入面板时取消隐藏计时"""
        if hasattr(self.camera_widget, '_hide_timer'):
            self.camera_widget._hide_timer.stop()
        super().enterEvent(event)

    def leaveEvent(self, event):
        """鼠标离开面板时启动隐藏计时"""
        if hasattr(self.camera_widget, '_hide_timer'):
            self.camera_widget._hide_timer.start()
        super().leaveEvent(event)


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("荣宇AI智慧分析系统")
        self.setMinimumSize(1200, 750)

        # 中央容器
        central = QWidget()
        main_layout = QVBoxLayout()
        main_layout.setContentsMargins(12, 8, 12, 8)
        main_layout.setSpacing(8)

        # 顶部标题栏
        header = QHBoxLayout()
        lbl_title = QLabel("荣宇AI智慧分析系统")
        lbl_title.setFont(QFont("Microsoft YaHei", 18, QFont.Weight.Bold))
        lbl_title.setStyleSheet("color: #4a9eff; background: transparent; padding: 4px 0;")
        header.addWidget(lbl_title)
        header.addStretch()
        main_layout.addLayout(header)

        # 分隔线
        line = QFrame()
        line.setFrameShape(QFrame.Shape.HLine)
        line.setStyleSheet("background-color: #333; max-height: 1px;")
        main_layout.addWidget(line)

        # 左右分栏（可拖拽）
        self.splitter = QSplitter(Qt.Orientation.Horizontal)
        self.splitter.setHandleWidth(6)
        self.splitter.setStyleSheet("""
            QSplitter::handle {
                background-color: #333;
                border-radius: 2px;
            }
            QSplitter::handle:hover {
                background-color: #4a9eff;
            }
        """)

        # 左侧：相机画面
        self.camera_widget = CameraWidget()
        self.camera_widget.setMinimumWidth(400)

        # 右侧：分析面板（Tab 切换藻类/虾分析）
        self.right_tabs = QTabWidget()
        self.right_tabs.setMinimumWidth(280)
        self.right_tabs.setStyleSheet("""
            QTabWidget::pane {
                border: 1px solid #333;
                border-radius: 4px;
                background-color: #16213e;
            }
            QTabBar::tab {
                background-color: #1a1a2e;
                color: #999;
                padding: 8px 16px;
                font-size: 13px;
                font-weight: bold;
                border: 1px solid #333;
                border-bottom: none;
                border-top-left-radius: 6px;
                border-top-right-radius: 6px;
                margin-right: 2px;
            }
            QTabBar::tab:selected {
                background-color: #16213e;
                color: #4a9eff;
                border-bottom: 2px solid #4a9eff;
            }
            QTabBar::tab:hover:!selected {
                background-color: #222244;
                color: #ccc;
            }
        """)

        self.algae_panel = AnalysisPanel(self.camera_widget, mode="algae")
        self.shrimp_panel = AnalysisPanel(self.camera_widget, mode="shrimp")

        self.right_tabs.addTab(self.algae_panel, "藻类分析")
        self.right_tabs.addTab(self.shrimp_panel, "虾体分析")

        self.splitter.addWidget(self.camera_widget)
        self.splitter.addWidget(self.right_tabs)
        self.splitter.setStretchFactor(0, 3)
        self.splitter.setStretchFactor(1, 1)
        self.splitter.setSizes([900, 300])

        main_layout.addWidget(self.splitter, 1)
        central.setLayout(main_layout)
        self.setCentralWidget(central)

        # 全局样式
        self.setStyleSheet("""
            QMainWindow { background-color: #0f0f1a; }
            QWidget { background-color: #16213e; color: #e0e0e0; }
            QPushButton {
                background-color: #2a2a4a;
                color: #e0e0e0;
                border: 1px solid #444;
                border-radius: 4px;
                padding: 6px 16px;
                font-size: 13px;
            }
            QPushButton:hover { background-color: #3a3a5a; }
        """)

    def closeEvent(self, event):
        self.camera_widget.closeCamera()
        event.accept()


def main():
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec())


if __name__ == '__main__':
    main()
